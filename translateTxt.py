import logging
import os, re
from docx import Document
from translator import TextTranslator
from translate import GoogleTranslateScraper


logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")


def split_text_into_chapter_files(text, output_dir, conter= 20):
    os.makedirs(output_dir, exist_ok=True)  # Asegura que el directorio de salida exista
    
    chapters = []
    current_chapter = []
    
    lines = text.split("\n")  # Divide el texto por líneas
    for line in lines:
        if re.match(r"\s*Capítulo\s+\d+", line, re.IGNORECASE):
            if current_chapter:
                chapters.append("\n".join(current_chapter))  # Junta líneas y conserva saltos de línea
                current_chapter = []
        current_chapter.append(line)  # Agrega la línea actual al capítulo en curso
    
    if current_chapter:
        chapters.append("\n".join(current_chapter))  # Agregar el último capítulo
    
    for i in range(1, len(chapters), conter):
        block = chapters[i:i + conter] 
            
        block_start = i 
        block_end = block_start + len(block) - 1
        output_file = os.path.join(output_dir, f"capitulos_{block_start}_a_{block_end}.txt")
        
        with open(output_file, "w", encoding="utf-8") as output:
            output.write("\n\n".join(block))  # Une los capítulos con doble salto de línea
        
        logging.info(f"Archivo creado: {output_file}")


def split_text_by_number_chapter(text, output_dir, conter=20):
    """
    Split text in chunk chapters with begining number
    """
    os.makedirs(output_dir, exist_ok=True)  # Asegura que el directorio de salida exista

    chapters = []
    current_chapter = []
    expected_chapter = 1

    lines = text.split("\n")  # Divide el texto por líneas
    for line in lines:
        line_stripped = line.strip()
        # Detecta línea que es exactamente el capítulo esperado (ej. "1.", "2.", etc.)
        if re.match(rf"^\s*{expected_chapter}\.\s*", line_stripped):

            # Si ya había contenido, guardamos el capítulo anterior
            if current_chapter:
                chapters.append("\n".join(current_chapter))
                current_chapter = []
            # Reescribe la línea como "Capítulo X"
            current_chapter.append(f"Capítulo {line}")
            expected_chapter += 1
        else:
            current_chapter.append(line)

                

    if current_chapter:
        chapters.append("\n".join(current_chapter))  # Agrega el último capítulo

    for i in range(0, len(chapters), conter):
        block = chapters[i:i + conter] 
            

        block_start = i + 1  # Ajuste para que empiece desde 1
        block_end = block_start + len(block) - 1
        output_file = os.path.join(output_dir, f"capitulos_{block_start}_a_{block_end}.txt")

        with open(output_file, "w", encoding="utf-8") as output:
            output.write("\n\n".join(block))  # Une los capítulos con doble salto de línea

    logging.info(f"Archivo creado: {output_file}")


def delete_document_if_exists(filename):
    if os.path.exists(filename):
        os.remove(filename)
        return True
    return False

def save_text_to_docx(text, filename="output.docx"):
    if not text:
        logging.warning("El texto está vacío. No se creará el documento.")
        return False

    if delete_document_if_exists(filename):
        logging.info(f"Se sobrescribió el archivo {filename}")

    document = Document()
    for line in text.split("\n"):
        if line.startswith("Capítulo") or line.startswith("Chapter"):
            document.add_heading(line, level=2)
        else:
            document.add_paragraph(line)
    document.save(filename)
    logging.info(f"Documento guardado como {filename}")
    return True

def extract_text_from_docx(filename):
    try:
        document = Document(filename)
        text = "\n".join([paragraph.text for paragraph in document.paragraphs])
        return text
    except Exception as e:
        print(f"Error al leer el archivo {filename}: {e}")
        return ""



def read_txt_file(file_path):
    # Número de caracteres que deseas leer (ajusta según lo que necesites)
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        try:
            with open(file_path, 'r', encoding='gbk') as f:
                return f.read()
        except Exception as e:
            logging.info(f"Error al leer el archivo con GBK: {e}")
            return None


def save_text_to_file(filename: str, content: str):

    filepath = f"{filename}.txt"
    
    try:
        with open(filepath, "w", encoding="utf-8") as file:
            file.write(content)
        logging.info(f"Archivo '{filepath}' creado con éxito.")
    except Exception as e:
        logging.error(f"Error al crear el archivo: {e}")


def split_near_middle(text):
    if '\n' not in text:
        return text, '' 

    middle = len(text) // 2
    newline_indices = [i for i, char in enumerate(text) if char == '\n']
    closest = min(newline_indices, key=lambda i: abs(i - middle))

    return text[:closest], text[closest + 1:]  


def split_text_to_two_docx_files(text, name_docx):
    part1, part2 = split_near_middle(text)
    save_text_to_docx(part1, f"{name_docx}1.docx")
    save_text_to_docx(part2, f"{name_docx}2.docx")
    

def merge_docx_and_save_as_txt(name_docx):
    docx_text1 = extract_text_from_docx(f"{name_docx}1.docx")
    docx_text2 = extract_text_from_docx(f"{name_docx}2.docx")
    docx_text = docx_text1 + docx_text2
    save_text_to_file(f"{name_docx}", docx_text)
 

def translate_text_and_save_txt(text,output_name):
    translator = TextTranslator()
    translated_text = translator.translate(text)
    save_text_to_file(output_name, translated_text)
    
def delete_posdata_line(text):
    lines = text.splitlines()
    new_line = [line for line in lines if not line.strip().lower().startswith("ps1:")]
    return "\n".join(new_line)

def add_dot_to_chapter_lines(text):
    lines = text.splitlines()
    updated_lines = []
    for line in lines:
        if line.strip().lower().startswith("capítulo") and not line.strip().endswith("."):
            updated_lines.append(line.rstrip() + ".")
        else:
            updated_lines.append(line)
    return "\n".join(updated_lines)

def delete_title_chapter(text):
    lines = text.splitlines()
    updated_lines = []
    conter_chapter = 0
    for line in lines:
        if line.strip().lower().startswith("capítulo") or line.strip().lower().startswith("el maestro desvergonzado de naruto capítulo"):
            conter_chapter += 1
            updated_lines.append(f"Capítulo {conter_chapter}.")
        else:
            updated_lines.append(line)
    return "\n".join(updated_lines)

def remove_section_until_next_chapter(text):
    lines = text.splitlines()
    result = []
    skip = False

    for line in lines:
        if not skip:
            if line.strip().startswith("******"):
                skip = True  # start skipping
            else:
                result.append(line)
        else:
            if line.strip().lower().startswith("capítulo"):
                result.append(line)
                skip = False  # stop skipping

    return "\n".join(result)


def translate_docx_with_selenium(name_docx):
    translator = GoogleTranslateScraper()
    translator.translate(f"{name_docx}1.docx")
    logging.info(f"Se ha creado en la carpeta el documento {name_docx}1.docx")
    translator.translate(f"{name_docx}2.docx")
    logging.info(f"Se ha creado en la carpeta el documento {name_docx}2.docx")

def merge_files(name_docx):
    text1= read_txt_file(f"{name_docx}1.txt")
    text2= read_txt_file(f"{name_docx}2.txt")
    text = text1 + text2
    save_text_to_file(name_docx, text)

def get_confirmation() -> bool:
    while True:
        response = input("¿Están todos los requisitos para continuar? (s/n): ").strip().lower()
        if response in ["s", "si", "y", "yes"]:
            return True
        elif response in ["n", "no"]:
            return False
        else:
            print("Por favor, responde con 's' o 'n'.")

def main():
    text = read_txt_file("反派竟是我自己？.txt")
    name_docx = "malo"
    split_text_to_two_docx_files(text, name_docx)
    # Translate docxs
    try:          
        translate_docx_with_selenium(name_docx)
    except Exception as e:
        print(f"Ocurrió un error inesperado: {e}")

    confirmation = get_confirmation()

    if(confirmation):
        merge_docx_and_save_as_txt(name_docx)
        text = read_txt_file(f"{name_docx}.txt")
        text = add_dot_to_chapter_lines(text)
        #translate_text_and_save_txt(text)
        #remove_section_until_next_chapter(text)
        #new_text = delete_posdata_line(text)
        #save_text_to_file(name_docx, text)
        split_text_into_chapter_files(text, f"{name_docx} 20 por 20", 20)
    print("Fallo inebitablemente =(")

def translate_novels():
    name_docx = "malo copy"
    text = read_txt_file(f"{name_docx}.txt")
    split_text_into_chapter_files(text, f"{name_docx} 20 por 20", 20)
    

def red_docx_novel(name_docx):
    docx_text = extract_text_from_docx(f"{name_docx}.docx")
    save_text_to_file(f"{name_docx}", docx_text)

if __name__ == "__main__":
    main()
    #translate_novels()