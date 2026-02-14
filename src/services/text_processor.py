import os
import re
import logging
import requests
from docx import Document
from urllib.parse import urlparse, urlunparse
from src.config.paths import OUTPUT_DIR, AUTO_SAVE_TXT
from src.utils.time_utils import filename_timestamp

class TextProcessor:
    """
    Clase para manipular, traducir y guardar texto en distintos formatos.
    Incluye utilidades para limpieza, normalización, fragmentación y persistencia de texto.
    """

    def __init__(self, text=""):
        self.text = text.strip()

    # ----------------------------
    # 🔹 LIMPIEZA Y FORMATEO
    # ----------------------------
    def clean_text(self, text=None):
        """
        Elimina saltos de línea, tabulaciones y espacios extra del texto.
        """
        text = text or self.text
        return text.replace("\n", " ").replace("\r", " ").replace("\t", " ").strip()

    def normalize_quotes(self, text=None):
        """
        Reemplaza comillas tipográficas (‘’, “”) por comillas estándar (' y ").
        """
        text = text or self.text
        replacements = {
            "“": '"', "”": '"', "‘": "'", "’": "'"
        }
        for original, replacement in replacements.items():
            text = text.replace(original, replacement)
        return text

    def get_base_url(self, url):
        parsed = urlparse(url)
        base_url = urlunparse(parsed._replace(query="", fragment=""))
        return base_url

    # ----------------------------
    # 🔹 DIVISIÓN DE TEXTO
    # ----------------------------
    def chunk_text(self, max_length=5000, text=None):
        """
        Divide el texto en fragmentos (chunks) de un máximo de `max_length` caracteres.
        Útil para procesar o traducir textos largos en partes manejables.
        """
        text = text or self.text
        sentences = text.split(". ")
        chunks, current_chunk = [], ""

        for sentence in sentences:
            if len(current_chunk) + len(sentence) + 2 <= max_length:
                current_chunk += sentence + ". "
            else:
                chunks.append(current_chunk.strip())
                current_chunk = sentence + ". "
        
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return chunks
    

    def extract_line(self, text:str, position:int=0):
        text = text or self.text
        sentences = text.splitlines()
        return sentences[position]
    
    def search_number(self, text):
        return re.search(r"第\s*(\d+(?:\.\d+)?)", text)

    # ----------------------------
    # 🔹 TRADUCCIÓN
    # ----------------------------
    def translate_text(self, target_language, source_language="auto", text=None):
        """
        Traduce el texto usando el endpoint público de Google Translate.
        Retorna el texto traducido o el original si hay error.
        """
        text = text or self.text
        url_translator = "https://translate.googleapis.com/translate_a/single"

        params = {
            "client": "gtx",
            "sl": source_language,
            "tl": target_language,
            "dt": "t",
            "q": text
        }

        try:
            response = requests.get(url_translator, params=params, timeout=10)
            response.raise_for_status()
            result = response.json()

            if result and isinstance(result, list) and len(result) > 0:
                return result[0][0][0]
            else:
                logging.warning("⚠️ Estructura inesperada en la respuesta de traducción.")
                return text
        except requests.exceptions.RequestException as e:
            logging.error(f"❌ Error de conexión con Google Translate: {e}")
            return text
        except Exception as e:
            logging.error(f"⚠️ Error inesperado en la traducción: {e}")
            return text

    # ----------------------------
    # 🔹 CREACIÓN Y LECTURA DE ARCHIVOS
    # ----------------------------


    def create_text_file(self, text, filename="texto.txt", encoding="utf-8"):
        """
        Crea un archivo de texto (.txt) con el contenido especificado.
        Si ya existe, lo sobrescribe.
        """
        try:
            if not text.strip():
                logging.warning("⚠️ El texto está vacío. No se creará el documento.")
                return False
            
            file_path = os.path.join(OUTPUT_DIR, filename)

            if self.delete_if_exists(file_path):
                logging.info(f"Se sobrescribió el archivo existente: {filename}")

            with open(file_path, "w", encoding=encoding) as file:
                file.write(text)
            logging.info(f"✅ Archivo de texto creado exitosamente: {filename}")
            return True
        except Exception as e:
            logging.error(f"❌ Error al crear el archivo de texto: {e}")
            return False
        
    def create_autosave_file(self, text, encoding="utf-8"):
        try:
            file_path = os.path.join(AUTO_SAVE_TXT, f"autosave[{filename_timestamp()}].txt")
            with open(file_path, "w", encoding=encoding) as file:
                file.write(text)
        except Exception as e:
            logging.error(f"❌ Error al crear el archivo de autoguardado: {e}")
        
    def read_txt_file(self, file_path):
        """ 
        Extrae el contenido de un archivo de texto (.txt) y lo devuelve como string.
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='gbk') as f:
                    return f.read()
            except Exception as e:
                logging.info(f"Error al leer el archivo con GBK: {e}")
                return None

    def delete_if_exists(self, file_path):
        """
        Elimina un archivo existente antes de crear uno nuevo.
        Retorna True si lo eliminó, False si no existía.
        """
        if os.path.exists(file_path):
            os.remove(file_path)
            return True
        return False

    def create_docx_from_text(self, text, filename="output.docx"):
        """
        Crea un archivo DOCX con el texto dado.
        Los títulos que comienzan con 'Capítulo' o 'Chapter' se formatean como encabezados.
        """
        try:
            if not text.strip():
                logging.warning("⚠️ El texto está vacío. No se creará el documento.")
                return False
            
            file_path = os.path.join(OUTPUT_DIR, filename)

            if self.delete_if_exists(file_path):
                logging.info(f"Se sobrescribió el archivo existente: {filename}")

            document = Document()
            for line in text.split("\n"):
                if line.startswith("Capítulo") or line.startswith("Chapter"):
                    document.add_heading(line, level=2)
                else:
                    document.add_paragraph(line)
            document.save(file_path)
            logging.info(f"✅ Documento guardado como {filename}")
            return True
        except Exception as e:
            logging.error(f"❌ Error al crear el documento DOCX: {e}")
            return False

    def extract_text_from_docx(self, filename):
        """
        Extrae el texto de un archivo DOCX y lo devuelve como string.
        """
        try:
            file_path = os.path.join(OUTPUT_DIR, filename)
            document = Document(file_path)
            text = "\n".join([p.text for p in document.paragraphs])
            return text
        except Exception as e:
            logging.error(f"❌ Error al leer el archivo {filename}: {e}")
            return ""
