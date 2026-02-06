from translator import TextTranslator
import os
import time
import random
import requests
from bs4 import BeautifulSoup
from docx import Document
import logging
from fake_useragent import UserAgent

# Configurar el registro
logging.basicConfig(level=logging.INFO, format="%(asctime)s - %(levelname)s - %(message)s")

class WebScraper:
    def __init__(self, url, chapter_tag, chapter_class, content_tag, content_class, title_tag=None, title_class=None, title_id=None, content_id=None):
        self.url = url
        self.base_url = self.get_base_url(self.url)
        self.response = None
        self.soup = None
        self.chapter_tag = chapter_tag
        self.chapter_class = chapter_class
        self.content_tag = content_tag
        self.content_class = content_class
        self.title_tag = title_tag
        self.title_class = title_class
        self.content_id = content_id
        self.title_id = title_id

    def get_base_url(self, url):
        from urllib.parse import urlparse
        parsed_url = urlparse(url)
        return f"{parsed_url.scheme}://{parsed_url.netloc}"

    def fetch_page(self):
        """Obtiene el contenido de la página."""
        try:
            session = requests.Session()
            session.headers.update({
                "User-Agent": UserAgent().random,
                "Referer": "https://www.ddxs.com/",  
                "Accept-Language": "en-US,en;q=0.9",
            })

            response = session.get(self.url)
            print(response.status_code)
            self.soup = BeautifulSoup(response.text, 'html.parser')

            time.sleep(random.uniform(2, 5)) 
        except requests.exceptions.RequestException as err:
            logging.error(f"Error al obtener la página: {err}")
            return False
        return True


    def extract_links(self):
        """Extrae enlaces a los capítulos."""
        if not self.soup:
            logging.error("La página de los links no ha sido cargada.")
            return []
        elements = self.soup.find_all(self.chapter_tag, class_=self.chapter_class)
        links = []
        for element in elements:
            link = element.find('a')
            if link and 'href' in link.attrs:
                href = link['href']
                if href.startswith('/'):
                    href = f"{self.base_url}{href}"
                links.append(href)
        return links

    def extract_pagination_links(self):
        if not self.soup:
            logging.error("La página de los links no ha sido cargada.")
            return []

        container = self.soup.find(self.chapter_tag, class_=self.chapter_class)
        if not container:
            logging.warning("No se encontró el contenedor con los enlaces.")
            return []

        links = []
        for a in container.find_all('a', href=True):
            href = a['href']
            if href.startswith('/'):
                href = f"{self.base_url}{href}"
            links.append(href)

        return links

    def extract_text_by_class(self, tag, class_name=None):
        """Extrae texto basado en etiquetas y clases."""
        if not self.soup:
            logging.error("La página no ha sido cargada 2.")
            return ""
        element = self.soup.find(tag, class_=class_name)
        return element.text if element else ""
    

    def extract_text_by_id(self, tag, element_id=None):
        """Extrae texto basado en etiquetas y id."""
        if not self.soup:
            logging.error("La página no ha sido cargada.")
            return ""
        element = self.soup.find(tag, id=element_id)
        return element.text if element else ""
    

    def get_extract_title(self, title_tag=None, title_id=None, title_class=None):
        title_tag = title_tag if title_tag else self.title_tag
        title_id = title_id if title_id else self.title_id
        title_class = title_class if title_class else self.title_class
        if title_id:
            return self.extract_text_by_id(title_tag, title_id)
        elif title_class:
            return self.extract_text_by_class(title_tag, title_class)
        else:
            return None
    
    def get_extract_content(self, content_tag=None, content_id=None, content_class=None):
        content_tag = content_tag if content_tag else self.content_tag
        content_id = content_id if content_id else self.content_id
        content_class = content_class if content_class else self.content_class
        if content_id:
            return self.extract_text_by_id(content_tag, content_id)
        elif content_class :
            return self.extract_text_by_class(content_tag, content_class)
        else:
            return None


    def extract_chapters(self, links, max_chapters=20, initial_chapter=0):
        all_text = ""

        for i, link in enumerate(links[:max_chapters], start=1):
            if i < initial_chapter:
                logging.info(f"Ignorando el capítulo {i}")
                continue

            logging.info(f"Procesando capítulo {i} de {max_chapters}")
            self.url = link
            if not self.fetch_page():
                logging.warning(f"No se pudo abrir el URL del capítulo: {link}")
                continue

            # Extraer título
            title = self.get_extract_title()
            title = title if title else f"Capítulo {i}"
            all_text += f"{title}\n"

            # Extraer contenido
            content = self.get_extract_content()
            if content:
                all_text += content + "\n"
            else:
                logging.warning(f"No se encontró contenido en el capítulo {i}")

            # Pausa para evitar bloqueo
            if i % 10 == 0:
                logging.info("Pausa larga para evitar bloqueo...")
                time.sleep(random.uniform(10, 20))
        return all_text
    
    def extract_novels_title(self, tag, class_name=None):
        if not self.soup:
            logging.error("La página no ha sido cargada.")
            return []
        elements = self.soup.find_all(tag, class_=class_name)
        return [element.text for element in elements]


def delet_document_if_exists(filename):
    if os.path.exists(filename):
        os.remove(filename)
        return True
    return False


def create_document_from_text(text, filename="output.docx"):
    if not text:
        logging.warning("El texto está vacío. No se creará el documento.")
        return False

    if delet_document_if_exists(filename):
        logging.info(f"Se sobrescribió el archivo {filename}")

    document = Document()
    for line in text.split("\n"):
        if line.startswith("Capítulo") or line.startswith("Chapter"):
            document.add_heading(line, level=2)
        else:
            document.add_paragraph(line)
    document.save(filename)
    logging.info(f"Documento guardado como {filename}")
    return True


def create_text_file(text, filename="texto.txt", encoding="utf-8"):
    try:
        if not text:
            logging.warning("El texto está vacío. No se creará el documento.")
            return False

        if delet_document_if_exists(filename):
            logging.info(f"Se sobrescribió el archivo {filename}")

        with open(filename, "w", encoding=encoding) as file:
            file.write(text)
        logging.info(f"Archivo creado exitosamente: {filename}")
        return True
    except Exception as e:
        logging.warning(f"Error al crear el archivo: {e}")
        return False


def main():
    # Configuración específica para la página
    url = "https://jpxs123.com/tongren/index_54.html"
    scraper = WebScraper(
        url             = url,
        chapter_tag     = "div",
        chapter_class   = "page",
        title_tag       = None,
        title_class     = None,
        content_tag     = "div",
        content_class   = "infos",
        title_id        = None,
        content_id      = None
    )

    # Obtener enlaces de capítulos
    if scraper.fetch_page():
        links = scraper.extract_pagination_links()
        if not links:
            logging.error("No se encontraron enlaces de capítulos.")
        else:
            all_text = ""
            for e, link in enumerate(links[:20], start=1):
                all_text += "link: " + link + "\n"
                scraper.url = link
                if not scraper.fetch_page():
                    logging.warning(f"No se pudo abrir el URL del capítulo: {link}")
                else:
                    novels = scraper.extract_novels_title("h3", None)
                    if not novels:
                        logging.error("No se encontraron enlaces de capítulos.")
                    else:
                        initial = 0
                        for i, novel in enumerate(novels[:50], start=1):
                            if i < initial:
                                logging.info(f"Ignorando el capítulo {i}")
                                continue

                            logging.info(f"Procesando capítulo {i} de {len(novels)}")
                            if novel:
                                all_text += novel + "\n"
                            else:
                                logging.warning(f"No se encontró contenido en el capítulo {i}")


                        
                        
            #Crear el txt
            if create_text_file(all_text, filename="novelas.txt"):
                logging.info("Programa finalizado con Exito")
            else:
                logging.error("Error al crear el Documento")
                        

if __name__ == "__main__":
    main()



""" 
text = get_text()
translator = TextTranslator()
translated_text = translator.translate(text)
create_text_file(translated_text) 
"""
""" 
Marvel aprovechado
url = "https://www.ddxs.com/manweilidewaiguawanjia/"
scraper = WebScraper(
    url             = url,
    chapter_tag     = "td",
    chapter_class   = "L",
    title_tag       = "h1",
    title_class     = None,
    content_tag     = "dd",
    content_class   = None,
    title_id        = None,
    content_id      = "contents"
) """
""" 
Royal Scan
url=url,
chapter_tag="tr",
chapter_class="chapter-row",
content_tag="div",
content_class="chapter-inner chapter-content",
title_tag="h2",
title_class="mbs_posts_title" 
"""
""" 
def get_text():
    file_path = "木叶：这个宇智波只想躺平.txt"

    # Número de caracteres que deseas leer (ajusta según lo que necesites)

    try:
        with open(file_path, "r", encoding="utf-8") as file:
            first_part = file.read(3000)
            return first_part
    except FileNotFoundError:
        logging.info(f"El archivo {file_path} no se encontró.")
    except Exception as e:
        logging.info(f"Error al leer el archivo: {e}")
"""

""" 
def create_text_file(text, filename="texto.txt", encoding="utf-8"):
    try:
        with open(filename, "w", encoding=encoding) as file:
            file.write(text)
        logging.info(f"Archivo creado exitosamente: {filename}")
    except Exception as e:
        logging.warning(f"Error al crear el archivo: {e}")
 """

""" all_text = scraper.extract_chapters(links, initial_chapter=1, max_chapters=30 ) """

""" translator = TextTranslator()
translated_text = translator.translate(all_text) """


# Crear el documento
""" if create_document_from_text(translated_text, filename="Marvel algo 1 al 30.docx"):
    logging.info("Programa finalizado con Exito")
else:
    logging.error("Error al crear el Documento") """