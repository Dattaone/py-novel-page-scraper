import time
import random
import requests
from bs4 import BeautifulSoup
from docx import Document


class WebScraper:

    def __init__(self, url):
        self.url = url
        
        self.response = None
        self.soup = None
    
    def fetch_page(self):
        try:
            self.response = requests.get(self.url)
            self.response.raise_for_status
            self.soup = BeautifulSoup(self.response.text, 'html.parser')
            time.sleep(random.uniform(2,5))
        except requests.exceptions.RequestException as err:
            print(f"Error al obtener la pagina: {err}")
            return False
        return True
    
    def extract_chapters_links(self, tag, class_name=None):
        if not self.soup:
            print("Error: La página de los links no ha sido cargada")
            return []
        elements = self.soup.find_all(tag, class_=class_name)
        list_links = []
        for element in elements:
            link = element.find('a')
            if link and 'href' in link.attrs:
                href = link['href']
                if href.startswith('/'):
                    href = f"https://www.royalroad.com{href}"    
                list_links.append(href)

        return list_links
    
    def extract_chapter_text(self, tag, class_name=None):
        if not self.soup:
            print("Error: La página no ha sido cargada.")
            return ""
        elements = self.soup.find_all(tag, class_=class_name)
        chapter_text = ""
        for element in elements:
            chapter_text += element.text
        return chapter_text
    
    def translate_text(text, target_language, source_language="auto"):
        url_translator = "https://translate.googleapis.com/translate_a/single"

        params = {
            "client"    : "gtx",
            "sl"        : source_language,
            "tl"        : target_language,
            "dt"        : "t",
            "q"         : text
        }

        response = requests.get(url_translator, params=params)
        if response.status_code == 200:
            result = response.json()
            return result[0][0][0]
        else:
            raise Exception(f"Error en la traducción: {response.status_code}")


def clean_text(text):
    """Elimina caracteres problemáticos y espacios adicionales."""
    return text.replace("\n", " ").replace("\r", " ").replace("\t", " ").strip()

def normalize_quotes(text):
    """Reemplaza comillas inteligentes por comillas estándar."""
    replacements = {
        "“": '"',  # Comillas dobles de apertura
        "”": '"',  # Comillas dobles de cierre
        "‘": "'",  # Comillas simples de apertura
        "’": "'",  # Comillas simples de cierre
    }
    for original, replacement in replacements.items():
        text = text.replace(original, replacement)
    return text

def chunk_text(text, max_length=5000):
    sentences = text.split(". ")
    chunks = []
    current_chunk = ""

    for sentence in sentences:
        if len(current_chunk) + len(sentence) + 2 <= max_length:
            current_chunk += sentence + ". "
        else:
            chunks.append(current_chunk.strip())
            current_chunk = sentence + ". "
    
    if current_chunk.strip():
        chunks.append(current_chunk.strip())
    
    return chunks

if __name__ == "__main__":
    url = "https://www.royalroad.com/fiction/50328/the-hawkshaw-inheritance"
    scraper = WebScraper(url)

    if(scraper.fetch_page()):
        links = scraper.extract_chapters_links("tr", "chapter-row")
        if not links:
            print("Error: La página de los links no ha sido cargada")
        else:
            conter = 0
            all_text = ""
            #Create document
            document = Document()
            for link in links:
                conter += 1

                print(f"conter : {conter}")

                if(conter % 10 == 0):
                    print("Pausa larga...")
                    time.sleep(random.uniform(10,20))

                #chapter title
                """ 
                scraper_title = WebScraper(link)
                if(scraper_title.fetch_page()):
                    title_chapter = scraper_title.extract_chapter_text("h2","mbs_posts_title")
                    document.add_heading(title_chapter, level=2)
                    all_text += title_chapter     
                """

                #chapter text
                scraper_chapter = WebScraper(link)
                if(scraper_chapter.fetch_page()):
                    text_chapter = scraper_chapter.extract_chapter_text("div","chapter-inner chapter-content")
                    document.add_heading(f"chapter {conter}:", level=2)
                    document.add_paragraph(text_chapter)
                    all_text += text_chapter
                    
                else:
                    print(f"no logro abrir el url del  capitulo: {conter}")

            if not all_text.strip():
                print("Error: No hay texto para traducir.")
            else:
                #one_chunk = chunk_text(all_text)[0]
                #one_chunk = clean_text(one_chunk)
                #print(f"fragmento: {one_chunk}")
                print(all_text)
                document.save("La herencia de Hawkshaw.docx")
                
                

                # Guardar el texto traducido
                """ with open("translated_text.txt", "w", encoding="utf-8") as file:
                    file.write(all_text)

                print("Texto traducido guardado en 'translated_text.txt'") """
