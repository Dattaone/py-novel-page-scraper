import os
import re
import logging
import requests
from docx import Document
from urllib.parse import urlparse, urlunparse

class TextProcessor:
    """
    Clase para manipular, traducir y guardar texto en distintos formatos.
    Incluye utilidades para limpieza, normalizaci√≥n, fragmentaci√≥n y persistencia de texto.
    """

    def __init__(self, text=""):
        self.text = text.strip()

    # ----------------------------
    # üîπ LIMPIEZA Y FORMATEO
    # ----------------------------
    def clean_text(self, text=None):
        """
        Elimina saltos de l√≠nea, tabulaciones y espacios extra del texto.
        """
        text = text or self.text
        return text.replace("\n", " ").replace("\r", " ").replace("\t", " ").strip()

    def normalize_quotes(self, text=None):
        """
        Reemplaza comillas tipogr√°ficas (‚Äò‚Äô, ‚Äú‚Äù) por comillas est√°ndar (' y ").
        """
        text = text or self.text
        replacements = {
            "‚Äú": '"', "‚Äù": '"', "‚Äò": "'", "‚Äô": "'"
        }
        for original, replacement in replacements.items():
            text = text.replace(original, replacement)
        return text

    def get_base_url(self, url):
        parsed = urlparse(url)
        base_url = urlunparse(parsed._replace(query="", fragment=""))
        return base_url

    # ----------------------------
    # üîπ DIVISI√ìN DE TEXTO
    # ----------------------------
    def chunk_text(self, max_length=5000, text=None):
        """
        Divide el texto en fragmentos (chunks) de un m√°ximo de `max_length` caracteres.
        √ötil para procesar o traducir textos largos en partes manejables.
        """
        text = text or self.text
        sentences = text.split(". ")
        chunks, current_chunk = [], ""

        for sentence in sentences:
            if len(current_chunk) + len(sentence) + 2 <= max_length:
                current_chunk += sentence + ". "
            else:
                chunks.append(current_chunk.strip())
                current_chunk = sentence + ". "
        
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        return chunks
    

    def extract_line(self, text:str, position:int=0):
        text = text or self.text
        sentences = text.splitlines()
        return sentences[position]
    
    def search_number(self, text):
        return re.search(r"Á¨¨\s*(\d+(?:\.\d+)?)", text)

    # ----------------------------
    # üîπ TRADUCCI√ìN
    # ----------------------------
    def translate_text(self, target_language, source_language="auto", text=None):
        """
        Traduce el texto usando el endpoint p√∫blico de Google Translate.
        Retorna el texto traducido o el original si hay error.
        """
        text = text or self.text
        url_translator = "https://translate.googleapis.com/translate_a/single"

        params = {
            "client": "gtx",
            "sl": source_language,
            "tl": target_language,
            "dt": "t",
            "q": text
        }

        try:
            response = requests.get(url_translator, params=params, timeout=10)
            response.raise_for_status()
            result = response.json()

            if result and isinstance(result, list) and len(result) > 0:
                return result[0][0][0]
            else:
                logging.warning("‚ö†Ô∏è Estructura inesperada en la respuesta de traducci√≥n.")
                return text
        except requests.exceptions.RequestException as e:
            logging.error(f"‚ùå Error de conexi√≥n con Google Translate: {e}")
            return text
        except Exception as e:
            logging.error(f"‚ö†Ô∏è Error inesperado en la traducci√≥n: {e}")
            return text

    # ----------------------------
    # üîπ CREACI√ìN Y LECTURA DE ARCHIVOS
    # ----------------------------
    def create_text_file(self, text, filename="texto.txt", encoding="utf-8"):
        """
        Crea un archivo de texto (.txt) con el contenido especificado.
        Si ya existe, lo sobrescribe.
        """
        try:
            if not text.strip():
                logging.warning("‚ö†Ô∏è El texto est√° vac√≠o. No se crear√° el documento.")
                return False

            if self.delete_if_exists(filename):
                logging.info(f"Se sobrescribi√≥ el archivo existente: {filename}")

            with open(filename, "w", encoding=encoding) as file:
                file.write(text)
            logging.info(f"‚úÖ Archivo de texto creado exitosamente: {filename}")
            return True
        except Exception as e:
            logging.error(f"‚ùå Error al crear el archivo de texto: {e}")
            return False
        
    def read_txt_file(self, file_path):
        """ 
        Extrae el contenido de un archivo de texto (.txt) y lo devuelve como string.
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            try:
                with open(file_path, 'r', encoding='gbk') as f:
                    return f.read()
            except Exception as e:
                logging.info(f"Error al leer el archivo con GBK: {e}")
                return None

    def delete_if_exists(self, filename):
        """
        Elimina un archivo existente antes de crear uno nuevo.
        Retorna True si lo elimin√≥, False si no exist√≠a.
        """
        if os.path.exists(filename):
            os.remove(filename)
            return True
        return False

    def create_docx_from_text(self, text, filename="output.docx"):
        """
        Crea un archivo DOCX con el texto dado.
        Los t√≠tulos que comienzan con 'Cap√≠tulo' o 'Chapter' se formatean como encabezados.
        """
        try:
            if not text.strip():
                logging.warning("‚ö†Ô∏è El texto est√° vac√≠o. No se crear√° el documento.")
                return False

            if self.delete_if_exists(filename):
                logging.info(f"Se sobrescribi√≥ el archivo existente: {filename}")

            document = Document()
            for line in text.split("\n"):
                if line.startswith("Cap√≠tulo") or line.startswith("Chapter"):
                    document.add_heading(line, level=2)
                else:
                    document.add_paragraph(line)
            document.save(filename)
            logging.info(f"‚úÖ Documento guardado como {filename}")
            return True
        except Exception as e:
            logging.error(f"‚ùå Error al crear el documento DOCX: {e}")
            return False

    def extract_text_from_docx(self, filename):
        """
        Extrae el texto de un archivo DOCX y lo devuelve como string.
        """
        try:
            document = Document(filename)
            text = "\n".join([p.text for p in document.paragraphs])
            return text
        except Exception as e:
            logging.error(f"‚ùå Error al leer el archivo {filename}: {e}")
            return ""
